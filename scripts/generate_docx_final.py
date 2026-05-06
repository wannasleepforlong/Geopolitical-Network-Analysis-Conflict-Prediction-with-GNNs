#!/usr/bin/env python3
"""
generate_docx_final.py
========================
Generates the final research paper with ACTUAL results from augmented data.
"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import parse_xml


def set_cell_shading(cell, hex_color):
    shading = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"',
        hex_color.lstrip('#')
    ))
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    run.font.size = Pt({1: 18, 2: 14, 3: 12}.get(level, 12))
    run.font.bold = True
    run.font.name = "Calibri"


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    if bold: run.font.bold = True
    if italic: run.font.italic = True
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "E2E8F0")
        for r in hdr[i].paragraphs[0].runs:
            r.font.bold = True; r.font.size = Pt(10); r.font.name = "Calibri"
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
            for r in cells[i].paragraphs[0].runs:
                r.font.size = Pt(10); r.font.name = "Calibri"
    doc.add_paragraph()


def build_doc():
    doc = Document()
    
    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("GeoPredict: A Graph Neural Network Framework for Temporal Conflict Forecasting in Global Geopolitical Networks")
    r.font.size = Pt(20); r.font.bold = True; r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A); r.font.name = "Calibri"
    
    doc.add_paragraph()
    d = doc.add_paragraph(); d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = d.add_run(datetime.now().strftime("%B %d, %Y")); r.font.italic = True; r.font.color.rgb = RGBColor(0x5B, 0x64, 0x74)
    doc.add_paragraph()
    
    # Abstract
    add_heading(doc, "Abstract", 1)
    abstract = ("Accurately forecasting interstate conflict remains a key challenge in international relations and computational social science. "
                "Traditional approaches struggle with the relational, non-Euclidean structure of global politics. We present GeoPredict, "
                "a graph neural network framework that transforms geopolitical event streams from GDELT into temporal graphs for conflict prediction. "
                "We implement TemporalGCN and TemporalGAT architectures and evaluate them against logistic regression, random forest, and gradient boosting baselines. "
                "Using an augmented temporal dataset spanning 84 months (60 synthetic, 24 real), our baselines achieve up to 82.3% AUC-ROC, "
                "while the GNN variants reach 53.6% AUC, revealing a significant performance gap that highlights the challenge of learning "
                "graph-structured dependencies from limited real-world conflict data.")
    add_para(doc, abstract)
    
    # 1 Introduction
    add_heading(doc, "1. Introduction", 1)
    add_heading(doc, "1.1 The Problem", 2)
    add_para(doc, ("Forecasting armed conflict has preoccupied scholars and policymakers for centuries. Classical statistical models treat countries as independent observations, "
                    "discarding the network topology of international relations. Modern machine learning pipelines assume Euclidean feature spaces, struggling with temporal non-stationarity."))
    
    add_heading(doc, "1.2 The Opportunity", 2)
    add_para(doc, ("Graph Neural Networks (GNNs) operate directly on graph-structured data, propagating information along edges to learn topology-conditioned representations. "
                    "Spatio-temporal GNNs further combine graph convolutions with recurrent temporal modules, suited for geopolitical forecasting where structural dependencies span months."))
    
    add_heading(doc, "1.3 Objective", 2)
    add_para(doc, ("This paper introduces GeoPredict with three contributions: (1) a reproducible temporal graph pipeline from GDELT 2.0; "
                    "(2) TemporalGCN and TemporalGAT architectures; (3) empirical evaluation against strong tabular baselines on augmented data."))
    
    # 2 Related Work
    add_heading(doc, "2. Related Work", 1)
    add_para(doc, ("Event datasets like GDELT and ICEWS encode billions of political events. Classical VAR models and network analysis established that position matters in IR, "
                    "but the shift to deep graph learning for conflict prediction remains nascent (Kipf et al., 2018; Velickovic et al., 2018)."))
    
    # 3 Methodology
    add_heading(doc, "3. System Architecture & Methodology", 1)
    add_heading(doc, "3.1 Data Engineering", 2)
    add_para(doc, ("GeoPredict ingests GDELT 2.0 events, filters to major-power FIPS countries, and constructs monthly conflict/cooperation networks with Goldstein scores and tone. "
                    "We augment the original 24-month dataset with 60 synthetically generated months based on fitted empirical distributions, yielding 84 total months."))
    
    add_heading(doc, "3.2 GNN Framework", 2)
    add_para(doc, ("TemporalGCN couples an LSTM encoder with symmetrically normalized Graph Convolution (Kipf & Welling, 2017). "
                    "TemporalGAT replaces convolution with multi-head self-attention (Velickovic et al., 2018). Both project final embeddings through an edge-prediction MLP."))
    
    add_heading(doc, "3.3 Augmentation Strategy", 2)
    add_para(doc, ("To address severe data scarcity, we generate 60 months of synthetic history by sampling from per-dyad empirical distributions fit on real GDELT data. "
                    "The procedure preserves regional conflict patterns (Middle East, South Asia hotspots) and temporal autocorrelation via AR(1) processes on Goldstein scores. "
                    "The augmented dataset spans January 2019 through December 2025."))
    
    add_heading(doc, "3.4 Training Pipeline", 2)
    add_para(doc, ("All experiments use seed=42, temporal window T=12, chronological split (70/10/20), AdamW optimization with weight decay 1e-5, "
                    "ReduceLROnPlateau scheduling, early stopping (patience=10), and gradient clipping. Threshold optimization maximizes validation F1."))
    
    # 4 Implementation
    add_heading(doc, "4. Implementation Details", 1)
    add_para(doc, ("Implemented in Python 3.11 with PyTorch 2.x. Models range from 18,977 (TemporalGCN) to 75,265 (TemporalGAT) parameters. "
                    "Training on CPU takes 6-28 seconds per model. The framework supports CPU and CUDA inference."))
    
    # 5 Experiments
    add_heading(doc, "5. Experiments & Evaluation", 1)
    add_heading(doc, "5.1 Setup", 2)
    add_para(doc, ("Dataset: 20 major-power countries, 84 monthly periods (60 synthetic + 24 real), ~26,189 valid dyads. Effective supervised split after windowing: 50 train / 7 val / 15 test. "
                    "Test set is exclusively the real period months (Oct 2024-Dec 2025)."))
    
    add_heading(doc, "5.2 Results", 2)
    add_para(doc, "Table 1 reports test-set performance on conflict escalation prediction across 3,620 test dyads, all computed at the optimal validation threshold.", italic=True)
    
    headers = ["Model", "AUC", "Macro-F1", "Precision", "Recall", "Accuracy"]
    rows = [
        ["Logistic Regression", "0.8231", "0.7462", "0.6816", "0.9207", "0.7517"],
        ["Random Forest", "0.7984", "0.5081", "0.8684", "0.1870", "0.5898"],
        ["Gradient Boosting", "0.7773", "0.6028", "0.7979", "0.3467", "0.6387"],
        ["TemporalGCN (Ours)", "0.5274", "0.5000", "0.5108", "0.1870", "0.5163"],
        ["TemporalGAT (Ours)", "0.5363", "0.5556", "0.5106", "0.3144", "0.5188"],
    ]
    add_table(doc, headers, rows)
    
    add_para(doc, ("On the augmented dataset, traditional baselines dramatically outperform the GNN variants. Logistic Regression achieves 82.3% AUC-ROC and 92.1% recall, "
                    "suggesting that the engineered edge features (conflict/cooperation counts, Goldstein scale, tone, imbalance ratio) carry strong predictive signal. "
                    "By contrast, TemporalGCN and TemporalGAT plateau at approximately 53% AUC—only marginally above random chance. Several factors explain this gap: "
                    "(1) Model capacity: the GNNs have 50K-75K parameters, creating severe overfitting on ~50 training windows. "
                    "(2) Distribution shift: the test set contains only real GDELT data, while 71% of training is synthetic. Synthetic data captures marginal distributions but misses "
                    "higher-order temporal autocorrelation and geopolitical shocks. "
                    "(3) Label construction: the Goldstein-drop label formulation (drop > 0.5 month-to-month) is noisy and leads to ~46% positives even in synthetic data, "
                    "making structural inductive biases less discriminative than tabular features. "
                    "(4) Message passing may dilute signal: with only 20 nodes, full-graph convolution averages away dyad-specific signals."))
    
    add_heading(doc, "5.3 Ablation: Effect of Dataset Size", 2)
    add_para(doc, ("We trained identical models on the original 24-month unaugmented dataset (8 train / 1 val / 3 test). The same baselines achieved 78.4% AUC-ROC, "
                    "demonstrating that data augmentation did improve baseline performance (+3.9 AUC points for Logistic Regression) but failed to lift GNNs beyond ~53% AUC. "
                    "This suggests the primary bottleneck is not sample count but the fundamental mismatch between GNN structural priors and the problem formulation."))
    
    # 6 Discussion
    add_heading(doc, "6. Discussion", 1)
    add_para(doc, ("The principal finding is unequivocal: for this geopolitical conflict prediction task, flat tabular models substantially outperform graph-based architectures. "
                    "This does not invalidate GNNs for IR but reveals important boundary conditions. GNNs excel when graph structure itself is predictive (e.g., alliance networks, trade dependencies); "
                    "when the target variable is primarily driven by dyadic feature values (event counts, sentiment scores), graph convolution may add noise through neighborhood averaging. "
                    "Additionally, with only 20 nodes, the information bottleneck from adjacency-based message passing is minimal, removing the GNN's comparative advantage over fully connected or flat models."))
    
    add_para(doc, ("Limitations include synthetic data fidelity (marginal distributions without geopolitical logic), the Goldstein-drop label heuristic (alternative formulations warrant exploration), "
                    "and the major-power whitelist excluding non-state actors and smaller states. Ethical considerations around predictive conflict models remain critical: "
                    "all results should be paired with human oversight, uncertainty quantification, and bias auditing."))
    
    # 7 Conclusion
    add_heading(doc, "7. Conclusion & Future Work", 1)
    add_para(doc, ("GeoPredict provides a reproducible pipeline for temporal conflict forecasting with GNNs. On our augmented 84-month dataset, traditional baselines achieve 82.3% AUC-ROC, "
                    "while TemporalGCN/TAT reach approximately 53%. Future directions include: (1) incorporating multimodal signals (UN voting, trade flows, military expenditure) as node attributes; "
                    "(2) alternative label formulations (e.g., UCDP conflict onset rather than Goldstein drops); (3) graph transformer architectures; "
                    "(4) real-time streaming with live GDELT feeds; and (5) causal counterfactual simulation via learned graph generative models."))
    
    # References
    add_heading(doc, "References", 1)
    refs = [
        "Kipf, T. N., & Welling, M. (2017). Semi-Supervised Classification with Graph Convolutional Networks. ICLR 2017.",
        "Velickovic, P., Cucurull, G., Casanova, A., Romero, A., Lio, P., & Bengio, Y. (2018). Graph Attention Networks. ICLR 2018.",
        "Leetaru, K., & Schrodt, P. A. (2013). GDELT: Global Data on Events, Language, and Tone, 1979-2012.",
        "Tetlock, P. E., & Gardner, D. (2015). Superforecasting: The Art and Science of Prediction.",
        "Maoz, Z. (2010). Networks of Nations: The Evolution, Structure, and Impact of International Networks, 1816-2001.",
    ]
    for r in refs:
        p = doc.add_paragraph(style="List Number")
        p.add_run(r).font.size = Pt(10)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(3)
    
    doc.add_paragraph()
    fp = doc.add_paragraph(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("(c) 2025 GeoPredict Research Group"); r.font.size = Pt(9); r.font.italic = True; r.font.color.rgb = RGBColor(0x5B, 0x64, 0x74)
    
    return doc


def main():
    out_path = Path("GeoPredict_Research_Paper.docx")
    doc = build_doc()
    doc.save(str(out_path))
    print(f"Saved to {out_path.resolve()}")


if __name__ == "__main__":
    main()
